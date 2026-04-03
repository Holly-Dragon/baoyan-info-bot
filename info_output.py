import os
import pandas as pd
from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import docx
from info_get import export_categorized_excel, format_date_str

def add_hyperlink(paragraph, url, text, color="0000FF", underline=True):
    """为段落添加可点击的超链接 (依赖 python-docx 底层 XML)"""
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
    hyperlink.set(docx.oxml.shared.qn('r:id'), r_id)

    new_run = docx.oxml.shared.OxmlElement('w:r')
    rPr = docx.oxml.shared.OxmlElement('w:rPr')

    # 字体颜色
    if color:
        c = docx.oxml.shared.OxmlElement('w:color')
        c.set(docx.oxml.shared.qn('w:val'), color)
        rPr.append(c)
    
    # 字体下划线
    if underline:
        u = docx.oxml.shared.OxmlElement('w:u')
        u.set(docx.oxml.shared.qn('w:val'), 'single')
        rPr.append(u)
        
    # 设置字体为微软雅黑，字号小五 (9pt => 18 half-pts)
    rFonts = docx.oxml.shared.OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), '微软雅黑')
    rFonts.set(qn('w:eastAsia'), '微软雅黑')
    rFonts.set(qn('w:hAnsi'), '微软雅黑')
    rPr.append(rFonts)
    
    sz = docx.oxml.shared.OxmlElement('w:sz')
    sz.set(qn('w:val'), str(9 * 2))
    rPr.append(sz)
    szCs = docx.oxml.shared.OxmlElement('w:szCs')
    szCs.set(qn('w:val'), str(9 * 2))
    rPr.append(szCs)

    new_run.append(rPr)
    
    text_elem = docx.oxml.shared.OxmlElement('w:t')
    text_elem.text = text
    new_run.append(text_elem)
    
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return hyperlink

def set_run_font(run, font_name, pt_size, bold=False, color=None):
    """设置字体，兼容中文字体设置"""
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run.font.size = Pt(pt_size)
    run.bold = bold
    if color:
        run.font.color.rgb = color

def export_excel_to_word(excel_file="2026院校信息_分类.xlsx", word_file="2026院校信息_汇总.docx"):
    """
    将给定excel按照格式写出到word中。
    1. 按照sheet分类提取：学校、招生项目、通知链接、申请截止时间
    2. 按照规定的格式要求生成 word 文件。
    """
    if Document is None:
        print("请运行 `pip install python-docx` 安装库后再导出 Word 文档。")
        return
        
    if not os.path.exists(excel_file):
        print(f"未找到文件 `{excel_file}`，无法生成 Word 文档。")
        return

    print(f"\n=== 开始导出 Word 文档 ===")
    doc = Document()
    
    dfs = pd.read_excel(excel_file, sheet_name=None)
    
    for sheet_name, df in dfs.items():
        if df.empty:
            continue
            
        # 1. 大类名称（sheet 名）
        p_category = doc.add_paragraph()
        p_category.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_cat = p_category.add_run(sheet_name)
        # 小三 = 15pt，加粗，红色，居中
        set_run_font(run_cat, '微软雅黑', 15, bold=True, color=RGBColor(255, 0, 0))
        
        # 按照院校名称进行合并
        df = df.fillna('')
        grouped = df.groupby('学校', sort=False)
        
        for school, group in grouped:
            if not school:
                continue
                
            # 2. 院校名称
            p_school = doc.add_paragraph()
            p_school.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run_sch = p_school.add_run(str(school))
            # 五号 = 10.5pt，加粗，蓝色，左对齐
            set_run_font(run_sch, '微软雅黑', 10.5, bold=True, color=RGBColor(0, 128, 255))
            
            # 遍历该院校下的各个招生项目
            for _, row in group.iterrows():
                project = str(row.get('招生项目', '')).strip()
                link = str(row.get('通知链接', '')).strip()
                deadline = str(row.get('申请截止时间', '')).strip()
                
                if project:
                    # 3. 招生项目
                    p_proj = doc.add_paragraph()
                    p_proj.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    run_proj = p_proj.add_run(project)
                    # 五号 = 10.5pt，不加粗，黑色，左对齐
                    set_run_font(run_proj, '微软雅黑', 10.5, bold=False, color=RGBColor(0, 0, 0))
                    
                if link and link != '无链接':
                    # 4. 通知链接
                    p_link = doc.add_paragraph()
                    p_link.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    # 字体微软雅黑，字号小五 (9pt)，不加粗，蓝色，左对齐，且链接可点击
                    add_hyperlink(p_link, link, link, color="#0070c0", underline=True)

                if deadline:
                    # 5. 申请截止时间
                    p_deadline = doc.add_paragraph()
                    p_deadline.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    run_dl = p_deadline.add_run(f"报名截止时间：{deadline}")
                    # 五号 = 10.5pt，不加粗，蓝色，左对齐
                    set_run_font(run_dl, '微软雅黑', 10.5, bold=False, color=RGBColor(0, 128, 255))
                
                # 每个项目写完后，留有一行空行
                doc.add_paragraph()

    doc.save(word_file)
    print(f"成功导出 Word 文档：{word_file}\n")

def merge_and_output(main_file="2026院校信息.xlsx", update_file="2026院校信息_更新.xlsx", output_categorized="2026院校信息_分类.xlsx"):
    """
    读取更新文件，将其合并写入到主文件中，并重新生成分类版 Excel。
    """
    if not os.path.exists(update_file):
        print(f"未找到更新文件 `{update_file}`，可能没有新数据。合并流程结束。")
        return
        
    print(f"\n=== 开始读取并合并更新数据 ===")
    
    all_data_map = {}
    
    # 读入主文件（如果存在）
    if os.path.exists(main_file):
        df_main = pd.read_excel(main_file)
        df_main = df_main.fillna('')
        
        # # 将已有的旧字段名替换成新字段名（处理旧表格包含的问题）
        # if '院校发布时间' in df_main.columns:
        #     df_main.rename(columns={'院校发布时间': '院校报名信息发布时间'}, inplace=True)
            
        main_data = df_main.to_dict('records')
        for idx, item in enumerate(main_data):
            # 兼容旧数据的日期格式化
            item['院校报名信息发布时间'] = format_date_str(item.get('院校报名信息发布时间', ''))
            item['申请截止时间'] = format_date_str(item.get('申请截止时间', ''))
            link = item.get('通知链接')
            if link and link != '无链接':
                all_data_map[link] = item
            else:
                # 确保无链接的历史数据不被覆盖丢失
                all_data_map[f"old_no_url_{idx}"] = item
            
    # 读入更新文件所有 sheets 并合并
    # read_excel 返回 dict of DataFrames 如果 sheet_name=None
    dfs_update = pd.read_excel(update_file, sheet_name=None)
    update_count = 0
    
    # 建立一个去重集合，避免同一个对象如果在_更新中有好几个分类被重复加入（如果是无链接项）
    seen_updates = set()
    
    for sheet_name, df_update in dfs_update.items():
        if df_update.empty:
            continue
        df_update = df_update.fillna('')
        update_data = df_update.to_dict('records')
        
        for idx, item in enumerate(update_data):
            link = item.get('通知链接')
            proj = item.get('招生项目')
            school = item.get('学校')
            uniq_key = f"{link}_{proj}_{school}"
            
            if uniq_key in seen_updates:
                continue
            seen_updates.add(uniq_key)
            update_count += 1
            
            if link and link != '无链接':
                all_data_map[link] = item
            else:
                all_data_map[f"new_no_url_{sheet_name}_{idx}"] = item
            
    print(f"读取到约 {update_count} 条更新（含交叉项）待合并。")
        
    final_list = list(all_data_map.values())
    print(f"合并后当前汇总数据共 {len(final_list)} 项，准备写入 Excel...")
    
    df_merged = pd.DataFrame(final_list)
    
    # 构建最终需要的字段及公式注入
    columns = ['更新时间', '院校报名信息发布时间', '学校', '学院', '专业', '招生项目', '通知链接', '申请截止时间', '申请倒计时', '是否截止']
    for col in columns:
        if col not in df_merged.columns:
            df_merged[col] = ''
    df_merged = df_merged[columns]
    
    # 确保根据格式化后的 'x年x月x日' 日期从早到晚排序
    df_merged['temp_date'] = pd.to_datetime(df_merged['院校报名信息发布时间'].str.extract(r'(\d+年\d+月\d+日)')[0], format='%Y年%m月%d日', errors='coerce')
    df_merged = df_merged.sort_values(by='temp_date', ascending=True, na_position='last').drop(columns=['temp_date']).reset_index(drop=True)
    
    # 注入截止到期日和状态的 Excel 函数公式（以H列 "x年x月x日" 作为解析源）
    row_indices = df_merged.index + 2
    df_merged['申请倒计时'] = [f'=IFERROR(IF(DATEVALUE(SUBSTITUTE(SUBSTITUTE(SUBSTITUTE(H{idx},"年","-"),"月","-"),"日",""))<TODAY(), "已截止", DATEVALUE(SUBSTITUTE(SUBSTITUTE(SUBSTITUTE(H{idx},"年","-"),"月","-"),"日",""))-TODAY() & "天"), "未知")' for idx in row_indices]
    df_merged['是否截止'] = [f'=IFERROR(IF(DATEVALUE(SUBSTITUTE(SUBSTITUTE(SUBSTITUTE(H{idx},"年","-"),"月","-"),"日",""))<TODAY(), "是", "否"), "未知")' for idx in row_indices]
    
    df_merged.to_excel(main_file, index=False)
    print(f"成功更新主文件：{main_file}")
    
    # 4. 生成分类文件
    export_categorized_excel(final_list, output_categorized)
    
    # 5. 生成 Word 汇总文档
    current_date_str = datetime.now().strftime("%m月%d日")
    export_excel_to_word(output_categorized, f"2026院校信息_{current_date_str}.docx")

if __name__ == "__main__":
    merge_and_output()
